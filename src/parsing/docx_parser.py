"""
DOCX Parser for lease documents
Extracts text, tables, and structure from Word documents
"""

import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.table import Table as DocxTable


@dataclass
class TableData:
    """Represents a parsed table from the document"""
    headers: List[str]
    rows: List[List[str]]
    raw_text: str
    table_type: str = "unknown"  # "rent_schedule", "data_sheet", "general"


@dataclass
class ParsedDocument:
    """Represents a fully parsed lease document"""
    file_path: str
    file_name: str
    full_text: str
    paragraphs: List[str]
    tables: List[TableData]
    sections: Dict[str, str]  # article/section -> content
    data_sheet: Dict[str, Any]
    tenant_name: str = ""


class DocxParser:
    """Parser for DOCX lease documents"""

    # Patterns for identifying document structure
    ARTICLE_PATTERN = re.compile(r'^ARTICLE\s+([IVXLCDM]+)[:\s]+(.+)$', re.IGNORECASE)
    SECTION_PATTERN = re.compile(r'^(\d+\.\d+)\s+(.+)$')
    EXHIBIT_PATTERN = re.compile(r'^EXHIBIT\s+([A-Z](?:-\d+)?)[:\s]*(.*)$', re.IGNORECASE)
    DATA_SHEET_PATTERN = re.compile(r'DATA\s+SHEET', re.IGNORECASE)

    def __init__(self):
        pass

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a DOCX file and extract all relevant content

        Args:
            file_path: Path to the DOCX file

        Returns:
            ParsedDocument with extracted content
        """
        path = Path(file_path)
        doc = Document(file_path)

        # Extract all paragraphs
        paragraphs = self._extract_paragraphs(doc)

        # Extract tables
        tables = self._extract_tables(doc)

        # Build full text
        full_text = self._build_full_text(paragraphs, tables)

        # Parse document structure (articles, sections)
        sections = self._parse_sections(paragraphs)

        # Extract data sheet
        data_sheet = self._extract_data_sheet(paragraphs, tables)

        # Extract tenant name from filename or data sheet
        tenant_name = self._extract_tenant_name(path.name, data_sheet)

        return ParsedDocument(
            file_path=str(path),
            file_name=path.name,
            full_text=full_text,
            paragraphs=paragraphs,
            tables=tables,
            sections=sections,
            data_sheet=data_sheet,
            tenant_name=tenant_name
        )

    def _extract_paragraphs(self, doc: Document) -> List[str]:
        """Extract all paragraphs from the document"""
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return paragraphs

    def _extract_tables(self, doc: Document) -> List[TableData]:
        """Extract all tables from the document"""
        tables = []
        for table in doc.tables:
            table_data = self._parse_table(table)
            if table_data:
                tables.append(table_data)
        return tables

    def _parse_table(self, table: DocxTable) -> Optional[TableData]:
        """Parse a single table into structured data"""
        rows = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            if any(cell for cell in row_data):  # Skip empty rows
                rows.append(row_data)

        if not rows:
            return None

        # First row is usually headers
        headers = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []

        # Build raw text representation
        raw_text = self._table_to_text(headers, data_rows)

        # Determine table type
        table_type = self._classify_table(headers, data_rows)

        return TableData(
            headers=headers,
            rows=data_rows,
            raw_text=raw_text,
            table_type=table_type
        )

    def _table_to_text(self, headers: List[str], rows: List[List[str]]) -> str:
        """Convert table to readable text format"""
        lines = []
        if headers:
            lines.append(" | ".join(headers))
            lines.append("-" * 50)
        for row in rows:
            lines.append(" | ".join(row))
        return "\n".join(lines)

    def _classify_table(self, headers: List[str], rows: List[List[str]]) -> str:
        """Classify the type of table based on content"""
        header_text = " ".join(headers).lower()

        # Check for rent schedule
        rent_keywords = ["rent", "year", "annual", "monthly", "lease year"]
        if any(kw in header_text for kw in rent_keywords):
            return "rent_schedule"

        # Check for data sheet fields
        data_sheet_keywords = ["tenant", "landlord", "premises", "square feet"]
        if any(kw in header_text for kw in data_sheet_keywords):
            return "data_sheet"

        return "general"

    def _build_full_text(self, paragraphs: List[str], tables: List[TableData]) -> str:
        """Build full text representation of the document"""
        parts = []
        parts.extend(paragraphs)
        for table in tables:
            parts.append(f"\n[TABLE: {table.table_type}]\n{table.raw_text}\n")
        return "\n\n".join(parts)

    def _parse_sections(self, paragraphs: List[str]) -> Dict[str, str]:
        """Parse document into sections based on article/section headers"""
        sections = {}
        current_section = "preamble"
        current_content = []

        for para in paragraphs:
            # Check for article header
            article_match = self.ARTICLE_PATTERN.match(para)
            if article_match:
                # Save previous section
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                current_section = f"Article {article_match.group(1)}: {article_match.group(2)}"
                current_content = [para]
                continue

            # Check for exhibit header
            exhibit_match = self.EXHIBIT_PATTERN.match(para)
            if exhibit_match:
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                current_section = f"Exhibit {exhibit_match.group(1)}"
                if exhibit_match.group(2):
                    current_section += f": {exhibit_match.group(2)}"
                current_content = [para]
                continue

            current_content.append(para)

        # Save final section
        if current_content:
            sections[current_section] = "\n".join(current_content)

        return sections

    def _extract_data_sheet(self, paragraphs: List[str], tables: List[TableData]) -> Dict[str, Any]:
        """Extract key-value pairs from the data sheet section"""
        data_sheet = {}

        # Find data sheet section in paragraphs
        in_data_sheet = False
        for i, para in enumerate(paragraphs):
            if self.DATA_SHEET_PATTERN.search(para):
                in_data_sheet = True
                continue

            if in_data_sheet:
                # Look for key-value patterns
                # Pattern: "Key: Value" or "Key Value"
                kv_match = re.match(r'^([A-Za-z\s]+?):\s*(.+)$', para)
                if kv_match:
                    key = kv_match.group(1).strip().lower().replace(" ", "_")
                    value = kv_match.group(2).strip()
                    data_sheet[key] = value

                # Look for numbered paragraph patterns like "(12) Security Deposit: Section 3.5"
                # The actual value is often on the next line
                numbered_match = re.match(r'^\((\d+)\)\s*([^:]+):\s*(.+)$', para)
                if numbered_match:
                    key = numbered_match.group(2).strip().lower().replace(" ", "_")
                    ref_value = numbered_match.group(3).strip()

                    # Look ahead for the actual value (usually dollar amount or specific value)
                    for j in range(1, 5):
                        if i + j < len(paragraphs):
                            next_para = paragraphs[i + j].strip()
                            # Check if it looks like a value (starts with $, contains numbers, etc.)
                            if next_para and (
                                next_para.startswith('$')
                                or re.match(r'^[\$\d,\.]+', next_para)
                                or ('$' in next_para and any(c.isdigit() for c in next_para))
                            ):
                                data_sheet[key] = next_para
                                data_sheet[f"{key}_reference"] = ref_value
                                break
                            elif next_para and not next_para.startswith('('):
                                # Non-empty, non-numbered paragraph could be the value
                                if any(c.isdigit() for c in next_para):
                                    data_sheet[key] = next_para
                                    data_sheet[f"{key}_reference"] = ref_value
                                    break

                # Stop at next major section
                if self.ARTICLE_PATTERN.match(para) or self.EXHIBIT_PATTERN.match(para):
                    break

        # Also extract from data sheet tables
        for table in tables:
            if table.table_type == "data_sheet":
                for row in table.rows:
                    if len(row) >= 2:
                        key = row[0].strip().lower().replace(" ", "_")
                        value = row[1].strip()
                        if key and value:
                            data_sheet[key] = value

        return data_sheet

    def _extract_tenant_name(self, filename: str, data_sheet: Dict[str, Any]) -> str:
        """Extract tenant name from filename or data sheet"""
        # Try data sheet first
        for key in ["tenant", "tenant_name", "trade_name"]:
            if key in data_sheet:
                return data_sheet[key]

        # Extract from filename
        # Format: "Tenant Name - Location - Other Info.docx"
        name_part = filename.split(" - ")[0]
        # Clean up version numbers and extensions
        name_part = re.sub(r'\s*\([^)]+\)\s*', '', name_part)
        name_part = re.sub(r'\.docx?$', '', name_part, flags=re.IGNORECASE)

        return name_part.strip()


def parse_all_leases(lease_dir: str) -> List[ParsedDocument]:
    """
    Parse all lease documents in a directory

    Args:
        lease_dir: Path to directory containing lease documents

    Returns:
        List of ParsedDocument objects
    """
    parser = DocxParser()
    documents = []

    lease_path = Path(lease_dir)
    for file_path in lease_path.glob("*.docx"):
        try:
            doc = parser.parse(str(file_path))
            documents.append(doc)
        except Exception as e:
            print(f"Error parsing {file_path}: {e}")

    return documents
